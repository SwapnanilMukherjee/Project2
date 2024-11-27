# backend/editor/services/document_service.py
from typing import Optional, List, Dict, Any
from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction
import hashlib
import markdown
from docx import Document as DocxDocument

from ..models import Document, DocumentVersion, Change
from .merge_service import MergeService

class DocumentService:
    def __init__(self):
        self.merge_service = MergeService()

    def create_document(self, title: str, passkey: str) -> Document:
        """Create a new document with initial empty content"""
        # Hash the passkey
        hashed_passkey = self._hash_passkey(passkey)
        
        # Create document
        document = Document.objects.create(
            title=title,
            passkey=hashed_passkey
        )
        document.initialize_content()
        
        # Create initial version
        DocumentVersion.objects.create(
            document=document,
            version=1.0,
            content=document.content
        )
        
        return document

    def get_document(self, document_id: str, passkey: str) -> Optional[Document]:
        """Get document if passkey is valid"""
        try:
            document = Document.objects.get(id=document_id)
            if self._verify_passkey(passkey, document.passkey):
                return document
        except ObjectDoesNotExist:
            pass
        return None

    def apply_changes(self, document: Document, changes: List[Change]) -> bool:
        """Apply multiple changes to document"""
        with transaction.atomic():
            try:
                for change in changes:
                    success = self.merge_service.apply_change(document, change)
                    if not success:
                        transaction.set_rollback(True)
                        return False
                return True
            except Exception:
                transaction.set_rollback(True)
                return False

    def get_version_history(self, document: Document) -> List[Dict[str, Any]]:
        """Get document version history with changes"""
        versions = []
        for version in document.versions.all():
            changes = Change.objects.filter(
                document=document,
                source_version__gte=version.version - 0.1,
                source_version__lt=version.version
            ).order_by('timestamp')
            
            versions.append({
                'version': version.version,
                'timestamp': version.timestamp,
                'changes': [
                    {
                        'type': change.operation_type,
                        'position': change.position,
                        'content': change.content,
                        'attributes': change.attributes
                    }
                    for change in changes
                ]
            })
        
        return versions

    def restore_version(self, document: Document, version_number: float) -> bool:
        """Restore document to a previous version"""
        try:
            version = document.versions.get(version=version_number)
            with transaction.atomic():
                # Update document content
                document.content = version.content
                document.current_version += 0.1
                document.save()
                
                # Create new version record
                DocumentVersion.objects.create(
                    document=document,
                    version=document.current_version,
                    content=document.content
                )
                
                return True
        except ObjectDoesNotExist:
            return False

    def export_markdown(self, document: Document) -> str:
        """Export document as markdown"""
        piece_table = document.get_piece_table()
        content = piece_table.get_text()
        
        # Apply styles and formatting
        for style in piece_table.styles:
            # Convert style ranges to markdown syntax
            if 'bold' in style.attributes:
                content = self._apply_markdown_style(content, style, '**')
            if 'italic' in style.attributes:
                content = self._apply_markdown_style(content, style, '_')
                
        # Handle line formatting
        for line in piece_table.lines:
            if line.type == 'bullet':
                # Add bullet point markers
                content = self._add_line_prefix(content, line, '* ')
            elif line.type == 'heading':
                # Add heading markers
                level = line.properties.get('headingLevel', 1)
                content = self._add_line_prefix(content, line, '#' * level + ' ')
                
        return content

    def export_docx(self, document: Document) -> bytes:
        """Export document as Word document"""
        piece_table = document.get_piece_table()
        content = piece_table.get_text()
        
        docx = DocxDocument()
        current_paragraph = docx.add_paragraph()
        
        # Apply styles and formatting
        for style in piece_table.styles:
            start = style.start_offset
            end = start + style.length
            run = current_paragraph.add_run(content[start:end])
            
            if 'bold' in style.attributes:
                run.bold = True
            if 'italic' in style.attributes:
                run.italic = True
            if 'underline' in style.attributes:
                run.underline = True
            if 'fontSize' in style.attributes:
                run.font.size = style.attributes['fontSize']
                
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    def _hash_passkey(self, passkey: str) -> str:
        """Hash passkey using SHA-256"""
        return hashlib.sha256(passkey.encode()).hexdigest()

    def _verify_passkey(self, passkey: str, hashed_passkey: str) -> bool:
        """Verify if passkey matches hashed passkey"""
        return self._hash_passkey(passkey) == hashed_passkey

    def _apply_markdown_style(self, content: str, style: Any, marker: str) -> str:
        """Apply markdown style markers to content range"""
        start = style.start_offset
        end = start + style.length
        return content[:start] + marker + content[start:end] + marker + content[end:]

    def _add_line_prefix(self, content: str, line: Any, prefix: str) -> str:
        """Add prefix to line at specified position"""
        lines = content.split('\n')
        line_number = content[:line.offset].count('\n')
        if line_number < len(lines):
            lines[line_number] = prefix + lines[line_number]
        return '\n'.join(lines)